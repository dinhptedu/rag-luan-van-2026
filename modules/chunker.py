import re, unicodedata
from pathlib import Path
from docx import Document
from dataclasses import dataclass
from typing import Optional

@dataclass
class Chunk:
    text: str
    source: str
    group: str
    chunk_id: int
    strategy: str
    section: str = ''
    parent_id: Optional[int] = None

def clean_vi(text):
    text = unicodedata.normalize('NFC', text)
    text = re.sub(r'[ \t]+', ' ', text)
    return text.strip()

def read_docx(path, group):
    doc = Document(path)
    paras = []
    for p in doc.paragraphs:
        t = clean_vi(p.text)
        if len(t) > 15:
            paras.append({'text':t,'style':p.style.name,
                          'source':Path(path).name,'group':group})
    for tbl in doc.tables:
        for row in tbl.rows:
            rt = ' | '.join(c.text.strip() for c in row.cells if c.text.strip())
            if len(rt)>15: paras.append({'text':clean_vi(rt),'style':'Table',
                                         'source':Path(path).name,'group':group})
    return paras

DIEU_RE = re.compile(r'^\s*(Dieu\s+\d+|CHUONG\s+[IVXLC\d]+)',re.IGNORECASE)

def chunk_fixed(paras, size=500, overlap=100):
    chunks, idx = [], 0
    for p in paras:
        text = p['text']
        for i in range(0, len(text), size-overlap):
            seg = text[i:i+size].strip()
            if len(seg)>30:
                chunks.append(Chunk(text=seg,source=p['source'],group=p['group'],
                    chunk_id=idx,strategy='fixed'))
                idx += 1
    return chunks

def chunk_hierarchical(paras, parent_size=800, child_size=200):
    # Buoc 1: tao Parent chunks theo Dieu/Chuong
    chunks, idx = [], 0
    parents, buf, sec, src, grp = [], [], None, '', ''
    for p in paras:
        if DIEU_RE.match(p['text']):
            if buf: parents.append({'text':' '.join(buf),'section':sec or '',
                                    'source':src,'group':grp})
            buf=[p['text']]; sec=p['text'][:60]; src=p['source']; grp=p['group']
        else:
            buf.append(p['text'])
            if not src: src,grp=p['source'],p['group']
    if buf: parents.append({'text':' '.join(buf),'section':sec or '',
                            'source':src,'group':grp})
    # Buoc 2: tao Child chunks tu moi Parent
    for par in parents:
        pid = idx
        chunks.append(Chunk(text=par['text'][:parent_size],source=par['source'],
            group=par['group'],chunk_id=idx,strategy='hierarchical_parent',
            section=par['section']))
        idx += 1
        text = par['text']
        for i in range(0, len(text), child_size):
            seg = text[i:i+child_size].strip()
            if len(seg)>30:
                chunks.append(Chunk(text=seg,source=par['source'],
                    group=par['group'],chunk_id=idx,strategy='hierarchical_child',
                    section=par['section'],parent_id=pid))
                idx += 1
    return chunks

def get_chunks(paras, strategy_cfg):
    s = strategy_cfg.get('strategy','fixed')
    if s == 'fixed':
        return chunk_fixed(paras, strategy_cfg.get('size',500), strategy_cfg.get('overlap',100))
    elif s == 'hierarchical':
        return chunk_hierarchical(paras, strategy_cfg.get('parent',800), strategy_cfg.get('child',200))
    else:  # sentence / structure — tuong tu fixed nhung tach khac
        return chunk_fixed(paras, 400, 80)
